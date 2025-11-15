import os
import tempfile
import uuid
import joblib
import pdfplumber
import docx
import re
import pandas as pd
import numpy as np
import shap
from flask import Flask, request, jsonify, send_file, url_for
from flask_cors import CORS
from werkzeug.utils import secure_filename
from docx import Document
from pathlib import Path
import datetime

# ---------------- CONFIG ---------------- #
ALLOWED_EXTENSIONS = {"pdf", "docx"}
UPLOAD_FOLDER = tempfile.gettempdir()

app = Flask(__name__)
CORS(app)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# ---------------- PHASE 1: CLAUSE CLASSIFICATION ---------------- #
try:
    clf_pipeline = joblib.load("models/logistic_tfidf_pipeline.pkl")
    print("✅ Phase 1 Model loaded successfully.")
except FileNotFoundError:
    print("⚠️ Classification model not found.")
    clf_pipeline = None
except Exception as e:
    print("⚠️ Classification model load error:", e)
    clf_pipeline = None

# Class map
try:
    df = pd.read_csv("clean_legal_clauses.csv")
    df = df.dropna(subset=["clean_text", "clause_type"])
    class_map = dict(enumerate(df["clause_type"].astype("category").cat.categories))
    print("✅ Class map loaded.")
except Exception as e:
    print(f"⚠️ Class map error: {e}")
    class_map = {}

# ---------------- PHASE 3: RISK MODEL + SHAP ---------------- #
try:
    risk_pipeline = joblib.load("models/logistic_reg_risk.pkl")
    print("✅ Risk model loaded successfully.")
except FileNotFoundError:
    print("⚠️ Risk model not found.")
    risk_pipeline = None
except Exception as e:
    print("⚠️ Risk model load error:", e)
    risk_pipeline = None

# Handle pipeline or standalone LogisticRegression
if hasattr(risk_pipeline, "named_steps"):
    vectorizer = risk_pipeline.named_steps.get("tfidf", None)
    clf = risk_pipeline.named_steps.get("clf", None)
else:
    clf = risk_pipeline
    try:
        vectorizer = joblib.load("models/tfidf_vectorizer.pkl")
        print("✅ TF-IDF vectorizer loaded separately.")
    except FileNotFoundError:
        vectorizer = None
        print("⚠️ TF-IDF vectorizer not found. Risk analysis will be limited.")
    except Exception as e:
        vectorizer = None
        print("⚠️ TF-IDF vectorizer load error:", e)

# SHAP explainer
try:
    shap_explainer = joblib.load("models/shap_explainer.pkl")
    print("✅ SHAP explainer loaded.")
except Exception as e:
    shap_explainer = None
    print(f"⚠️ SHAP explainer not loaded: {e}")

# ---------------- UTILITIES ---------------- #
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(path):
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF error: {e}")
    return text

def extract_text_from_docx(path):
    try:
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"DOCX error: {e}")
        return ""

def split_into_clauses(text):
    text = re.sub(r"WEBSITE DESIGN AGREEMENT", r"0. WEBSITE DESIGN AGREEMENT", text, 1)
    raw_clauses = re.split(r"(\n\d{1,2}\. )", text)
    clauses = []
    preamble = raw_clauses[0].strip()
    if len(preamble) > 20:
        clauses.append(preamble)
    for i in range(1, len(raw_clauses), 2):
        if i + 1 < len(raw_clauses):
            full_clause = (raw_clauses[i] + raw_clauses[i + 1]).strip()
            if len(full_clause) > 20:
                clauses.append(full_clause)
    if not clauses:
        clauses = [c.strip() for c in text.split("\n\n") if len(c.strip()) > 20]
        if not clauses:
            clauses = [c.strip() for c in text.split(". ") if len(c.strip()) > 20]
    return clauses

# Map key terms to human-readable explanations (example)
WORD_RISK_MAP = {
    "assignment": "allows transfer of rights without restrictions",
    "ten": "contains ambiguous numeric thresholds",
    "business": "affects multiple business entities, increasing exposure",
    "party": "unclear responsibilities or obligations",
    "confidential": "lack of proper confidentiality clauses",
    # Add more mappings as needed
}

def generate_human_readable_justification(top_words):
    explanations = []
    for w, v in top_words:
        if w.lower() in WORD_RISK_MAP:
            direction = "increases" if v > 0 else "reduces"
            explanations.append(f"{WORD_RISK_MAP[w.lower()]} ({direction} risk)")
        else:
            direction = "increases" if v > 0 else "reduces"
            explanations.append(f"'{w}' ({direction} risk)")
    if not explanations:
        return "Clause risk is unclear from the text."
    return " ".join(explanations) + " Suggest clarifying or adding missing terms to reduce risk."

def analyze_risk_with_model(clause):
    """
    Predict risk and provide human-readable explainability
    """
    if not clf or not vectorizer:
        return {
            "risk_level": "Unknown",
            "confidence": 0.0,
            "justification": "Risk model or vectorizer not loaded."
        }

    vec = vectorizer.transform([clause])
    pred = clf.predict(vec)[0]
    prob = clf.predict_proba(vec).max()

    justification = "Explainability not available."
    if shap_explainer:
        try:
            shap_values = shap_explainer(vec)
            shap_vals_flat = np.array(shap_values.values).flatten()
            feature_importance = sorted(
                zip(vectorizer.get_feature_names_out(), shap_vals_flat),
                key=lambda x: abs(x[1]),
                reverse=True
            )[:5]

            # Generate human-readable justification
            justification = generate_human_readable_justification(feature_importance)
        except Exception as e:
            justification = f"Explainability error: {e}"

    return {
        "risk_level": pred,
        "confidence": float(prob) if hasattr(prob, "item") else float(prob),
        "justification": justification
    }

# ---------------- Assignment 2: Missing clause detection & modification ---------------- #

# Mandatory templates to append when missing (replace with legally-approved language if needed)
MANDATORY_CLAUSES = {
    "HIPAA": {
        "Data Privacy Protection Right":
            "Data Privacy Protection Right:\nThe Parties shall ensure that all protected health information and personal data processed under this Agreement "
            "are handled in accordance with applicable laws and standards (including HIPAA where applicable). "
            "Data subjects shall have appropriate rights to access, correction and to request deletion where applicable. "
            "Parties shall implement appropriate technical and organisational measures to protect personal data."
    },
    "GDPR": {
        "GDPR Data Protection Clause":
            "GDPR Data Protection Clause:\nThe Parties agree to comply with the EU General Data Protection Regulation (GDPR) where applicable. "
            "The Controller/Processor shall implement appropriate technical and organisational measures to ensure a level of security appropriate to the risk, "
            "adhere to data subject rights, and cooperate on breach notification and DPIA requirements."
    }
}

# Simple keyword sets used to decide whether clauses are present
DETECTION_KEYWORDS = {
    "hipaa": ["hipaa", "protected health information", "phi", "data privacy protection", "data privacy"],
    "gdpr": ["gdpr", "data subject", "data protection", "personal data", "data processing"]
}

def detect_regulatory_presence(text):
    """
    Quick keyword-based presence detection for HIPAA/GDPR related language.
    Returns a dict: {'hipaa_present': bool, 'gdpr_present': bool}
    """
    txt = text.lower()
    hipaa_present = any(k in txt for k in DETECTION_KEYWORDS["hipaa"])
    gdpr_present = any(k in txt for k in DETECTION_KEYWORDS["gdpr"])
    return {"hipaa_present": hipaa_present, "gdpr_present": gdpr_present}

def detect_missing_regulatory_clauses(full_text):
    """
    Determine which mandatory clauses are missing.
    Returns list of tuples (regime, clause_name)
    """
    presence = detect_regulatory_presence(full_text)
    missing = []

    # If neither HIPAA terms nor explicit 'Data Privacy Protection Right' phrase present, consider HIPAA clause missing
    if not presence["hipaa_present"] and "data privacy protection right" not in full_text.lower():
        missing.append(("HIPAA", "Data Privacy Protection Right"))

    # If GDPR not present or no GDPR clause text, consider GDPR missing
    if not presence["gdpr_present"] and "gdpr data protection clause" not in full_text.lower():
        missing.append(("GDPR", "GDPR Data Protection Clause"))

    return missing

def append_mandatory_clauses_to_text(full_text, missing_list):
    """
    Append missing mandatory clauses to contract text and return appended text.
    """
    appended = full_text.strip() + "\n\n" + "=== AUTO-ADDED COMPLIANCE CLAUSES ===\n"
    for regime, clause_name in missing_list:
        clause_text = MANDATORY_CLAUSES.get(regime, {}).get(clause_name, f"{clause_name}\n[Clause text not available]")
        appended += f"\n\n{clause_text}\n"
    return appended

def create_docx_from_text_and_clauses(full_text, missing_list, out_dir=None):
    """
    Create a DOCX file that contains original extracted text and the appended missing clauses.
    Returns path to saved file.
    """
    if out_dir is None:
        out_dir = app.config["UPLOAD_FOLDER"]
    base = f"contract_modified_{uuid.uuid4().hex[:8]}"
    out_path = os.path.join(out_dir, f"{base}.docx")

    doc = Document()
    doc.add_heading("Original Contract (extracted)", level=1)
    for line in full_text.splitlines():
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.add_page_break()
    doc.add_heading("Auto-added Compliance Clauses", level=1)
    doc.add_paragraph(f"Generated on {datetime.datetime.utcnow().isoformat()} UTC")
    for regime, clause_name in missing_list:
        clause_text = MANDATORY_CLAUSES.get(regime, {}).get(clause_name, f"{clause_name}\n[No template]")
        doc.add_heading(f"{clause_name} ({regime})", level=2)
        for p in clause_text.split("\n"):
            if p.strip():
                doc.add_paragraph(p.strip())

    doc.save(out_path)
    return out_path

# ---------------- ROUTES ---------------- #
@app.route("/", methods=["GET"])
def home():
    return "<h3>✅ Legal Clause Risk Analysis API is running. Use /upload to POST files. Use /modify_contract to get modified docx.</h3>"

@app.route("/upload", methods=["POST"])
def upload_file():
    """
    Upload a PDF or DOCX, extract clauses, classify + risk analyze each clause.
    Additionally detect missing HIPAA/GDPR clauses and produce modified DOCX (saved server-side).
    Response JSON includes:
      - total_clauses
      - analysis: list of clause objects
      - missing_clauses: list of missing tuples
      - modified_contract_download: endpoint URL (if missing clauses exist)
    """
    if "file" not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files["file"]
    if file.filename == "" or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file type"}), 400

    filename = secure_filename(file.filename)
    tmp_path = os.path.join(app.config["UPLOAD_FOLDER"], f"in_{uuid.uuid4().hex}_{filename}")
    file.save(tmp_path)

    try:
        # Extract text
        if filename.lower().endswith(".pdf"):
            text = extract_text_from_pdf(tmp_path)
        else:
            text = extract_text_from_docx(tmp_path)

        # Clean up uploaded temp
        try:
            os.remove(tmp_path)
        except Exception:
            pass

        if not text or not text.strip():
            return jsonify({"error": "No readable text found."}), 422

        # Clause splitting
        clauses = split_into_clauses(text)
        if not clauses:
            return jsonify({"error": "No clauses detected."}), 422

        # Per-clause analysis
        results = []
        for clause in clauses:
            # Phase 1: Classification
            phase1_pred = "N/A"
            phase1_conf = 0.0
            if clf_pipeline:
                try:
                    pred_code = clf_pipeline.predict([clause])[0]
                    # classifier might return numeric codes - attempt cast
                    try:
                        code_int = int(pred_code)
                        phase1_pred = class_map.get(code_int, str(pred_code))
                    except Exception:
                        phase1_pred = str(pred_code)
                    try:
                        phase1_conf = float(round(clf_pipeline.predict_proba([clause]).max(), 3))
                    except Exception:
                        phase1_conf = 0.0
                except Exception as e:
                    print("Phase1 prediction error:", e)

            # Phase 3: Risk
            risk_output = analyze_risk_with_model(clause)

            results.append({
                "clause": clause,
                "phase1": {
                    "predicted_clause_type": phase1_pred,
                    "confidence": phase1_conf
                },
                "phase3": risk_output
            })

        # Detect missing regulatory clauses and create modified DOCX if any missing
        missing = detect_missing_regulatory_clauses(text)
        modified_download_url = None
        modified_filename = None
        if missing:
            # create appended text docx and save
            modified_path = create_docx_from_text_and_clauses(text, missing, out_dir=app.config["UPLOAD_FOLDER"])
            modified_filename = os.path.basename(modified_path)
            # create a download URL endpoint (client will call /download_modified?filename=...)
            modified_download_url = url_for("download_modified", filename=modified_filename, _external=False)

        response = {
            "total_clauses": len(clauses),
            "analysis": results,
            "missing_clauses": missing,
            "modified_contract_download": modified_download_url,
            "modified_contract_filename": modified_filename
        }
        return jsonify(response)
    except Exception as e:
        print("Upload error:", e)
        return jsonify({"error": str(e)}), 500

@app.route("/download_modified", methods=["GET"])
def download_modified():
    """
    Query param: filename (exact filename returned previously)
    Example: /download_modified?filename=contract_modified_ab12cd34.docx
    """
    filename = request.args.get("filename")
    if not filename:
        return jsonify({"error": "filename query param required"}), 400
    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found"}), 404
    # send file as attachment
    return send_file(file_path, as_attachment=True, download_name=filename)

# ---------------- MAIN ---------------- #
if __name__ == "__main__":
    # ensure upload folder exists
    os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
    app.run(host="0.0.0.0", port=5000, debug=False)
